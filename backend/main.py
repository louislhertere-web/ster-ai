from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
import anthropic
import os
import json
import re
import base64
import requests
from dotenv import load_dotenv
from googleapiclient.discovery import build
from google.oauth2 import service_account
import io
from googleapiclient.http import MediaIoBaseDownload
import PyPDF2
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
import tempfile
from datetime import date

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '.env'))

app = FastAPI(title="Ster-AI API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

SCOPES = ['https://www.googleapis.com/auth/drive']
CREDENTIALS_FILE = os.path.join(os.path.dirname(__file__), 'credentials.json')
FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
RESEND_API_KEY = os.getenv("RESEND_API_KEY")

COMPETITIONS = ["N1", "N2", "N3"]

def get_drive_service():
    creds_b64 = os.getenv("GOOGLE_CREDENTIALS_BASE64")
    if creds_b64:
        creds_json = base64.b64decode(creds_b64).decode('utf-8')
        creds_dict = json.loads(creds_json)
        creds = service_account.Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    else:
        creds = service_account.Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
    return build('drive', 'v3', credentials=creds)

def get_or_create_folder(service, nom, parent_id):
    query = f"name='{nom}' and '{parent_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
    results = service.files().list(q=query, fields="files(id, name)").execute()
    fichiers = results.get('files', [])
    if fichiers:
        return fichiers[0]['id']
    metadata = {
        'name': nom,
        'mimeType': 'application/vnd.google-apps.folder',
        'parents': [parent_id]
    }
    folder = service.files().create(body=metadata, fields='id').execute()
    return folder['id']

def get_structure(service):
    structure = {}
    for comp in COMPETITIONS:
        comp_id = get_or_create_folder(service, comp, FOLDER_ID)
        structure[comp] = {'id': comp_id, 'journees': {}}
        query = f"'{comp_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = service.files().list(q=query, fields="files(id, name)").execute()
        for folder in results.get('files', []):
            structure[comp]['journees'][folder['name']] = folder['id']
    return structure

def lire_pdf(service, file_id):
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)
    reader = PyPDF2.PdfReader(fh)
    texte = ""
    for page in reader.pages:
        texte += page.extract_text()
    return texte

def extraire_nom_match(texte_pdf):
    prompt = f"""Extrait uniquement le nom du match depuis ce rapport de football, au format "Equipe A vs Equipe B".
Reponds UNIQUEMENT avec le nom du match, rien d'autre.
Si tu ne trouves pas, reponds "Match inconnu".

RAPPORT :
{texte_pdf[:1000]}"""
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=50,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text.strip()

def analyser_paire(id_match, rapport_arbitre, rapport_delegue):
    arbitre_manquant = rapport_arbitre == "Rapport arbitre non disponible"
    delegue_manquant = rapport_delegue == "Rapport delegue non disponible"

    if arbitre_manquant or delegue_manquant:
        manquant = "arbitre" if arbitre_manquant else "delegue"
        rapport_disponible = rapport_delegue if arbitre_manquant else rapport_arbitre
        nom_match = extraire_nom_match(rapport_disponible)
        return {
            "match": nom_match,
            "priorite": "gris",
            "motif": f"Rapport {manquant} non disponible",
            "action": f"Relancer l'officiel {manquant} pour obtenir son rapport"
        }

    prompt = f"""Tu es un assistant expert en discipline sportive pour la Direction des Competitions Nationales de la FFF (championnats N1, N2, N3).

Tu analyses deux rapports officiels d'un meme match : le rapport de l'arbitre central (RA_AC) et le rapport du delegue (RA_DP). Ton role est d'aider le responsable DCN a identifier rapidement les dossiers qui necessitent son attention.

ID Match : {id_match}

RAPPORT ARBITRE (RA_AC) :
{rapport_arbitre[:3000]}

RAPPORT DELEGUE (RA_DP) :
{rapport_delegue[:3000]}

CLASSIFICATION :

ROUGE — A traiter en priorite. Utilise ROUGE uniquement si le rapport mentionne explicitement l'un de ces faits :
- Violence physique (coup, agression, crachat) sur un joueur, officiel ou spectateur
- Propos discriminatoires ou racistes (origine, religion, orientation sexuelle)
- Utilisation d'articles pyrotechniques (fumigenes, feux de bengale, petards)
- Envahissement du terrain par des supporters ou dirigeants
- Interruption ou arret definitif du match pour incident grave
- Menaces ou voies de fait sur l'arbitre ou ses assistants
- Comportement violent d'un dirigeant ou educateur
- Jet de projectiles

VERT — Aucune action requise. Utilise VERT dans tous ces cas :
- Match sans incident notable
- Avertissements (cartons jaunes) simples, meme en grand nombre
- Toute difference de libelle ou de formulation sur le motif d'un carton jaune entre les deux rapports : c'est toujours VERT, jamais une contradiction significative
- Expulsion par double avertissement (double carton jaune) : VERT en precisant "Expulsion par double avertissement de [NOM DU JOUEUR]"
- Expulsion pour faute grossiere sans violence caracterisee
- Difference de nombre de cartons jaunes entre les deux rapports
- Difference dans l'identite des joueurs avertis

JAUNE — A verifier. Utilise JAUNE UNIQUEMENT dans ces deux cas strictement :
- Les deux rapports se contredisent sur le score final du match
- Les deux rapports se contredisent sur l'identite du joueur exclu directement (carton rouge direct, pas double jaune) pour faute grave ou violence

Toute autre situation doit etre classee VERT. En cas de doute, classe VERT.

Reponds UNIQUEMENT en JSON avec ce format exact :
{{
  "match": "Nom equipe domicile vs Nom equipe visiteur (date si disponible)",
  "priorite": "rouge" ou "jaune" ou "vert",
  "motif": "fait precis constate dans le rapport, en une phrase courte",
  "action": "ce que le responsable DCN doit faire concretement"
}}"""

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    texte = message.content[0].text
    json_match = re.search(r'\{.*\}', texte, re.DOTALL)
    return json.loads(json_match.group())

def envoyer_email(destinataire, resultats, competition=None, journee=None):
    rouge = [r for r in resultats if r['priorite'] == 'rouge']
    jaune = [r for r in resultats if r['priorite'] == 'jaune']
    vert = [r for r in resultats if r['priorite'] == 'vert']
    gris = [r for r in resultats if r['priorite'] == 'gris']

    contexte = f"{competition} - {journee}" if competition and journee else date.today().strftime("%d/%m/%Y")

    doc = Document()
    doc.add_heading('Ster-AI - Recapitulatif hebdomadaire', 0)
    doc.add_paragraph(f'Analyse automatique des rapports arbitres et delegues - {contexte}')
    doc.add_paragraph(f'Total : {len(rouge)} prioritaire(s) | {len(jaune)} a verifier | {len(vert)} RAS | {len(gris)} rapport(s) manquant(s)')
    doc.add_paragraph('')

    if rouge:
        doc.add_heading('PRIORITAIRES', level=1)
        for r in rouge:
            doc.add_heading(r['match'], level=2)
            doc.add_paragraph(f'Motif : {r["motif"]}')
            doc.add_paragraph(f'Action : {r["action"]}')
            doc.add_paragraph('')

    if jaune:
        doc.add_heading('A VERIFIER', level=1)
        for r in jaune:
            doc.add_heading(r['match'], level=2)
            doc.add_paragraph(f'Motif : {r["motif"]}')
            doc.add_paragraph(f'Action : {r["action"]}')
            doc.add_paragraph('')

    if gris:
        doc.add_heading('RAPPORTS MANQUANTS', level=1)
        for r in gris:
            doc.add_heading(r['match'], level=2)
            doc.add_paragraph(f'Motif : {r["motif"]}')
            doc.add_paragraph(f'Action : {r["action"]}')
            doc.add_paragraph('')

    if vert:
        doc.add_heading('RAS', level=1)
        for r in vert:
            doc.add_heading(r['match'], level=2)
            doc.add_paragraph(f'Motif : {r["motif"]}')
            doc.add_paragraph('')

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()

    with open(tmp.name, 'rb') as f:
        docx_b64 = base64.b64encode(f.read()).decode('utf-8')
    os.remove(tmp.name)

    html = f"""
    <html><body style="font-family: Arial; padding: 20px;">
    <h1 style="color: #1a1a2e;">Ster-AI - Recapitulatif {contexte}</h1>
    <p>Bonjour,</p>
    <p>Veuillez trouver ci-dessous le recapitulatif automatique de la journee du {date.today().strftime("%d/%m/%Y")}.</p>
    <p>Bien a vous,</p>
    <div style="display: flex; gap: 20px; margin: 20px 0;">
        <div style="background: #fff5f5; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 32px; font-weight: bold; color: #e53e3e;">{len(rouge)}</div>
            <div>PRIORITAIRES</div>
        </div>
        <div style="background: #fffff0; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 32px; font-weight: bold; color: #d69e2e;">{len(jaune)}</div>
            <div>A VERIFIER</div>
        </div>
        <div style="background: #f0fff4; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 32px; font-weight: bold; color: #38a169;">{len(vert)}</div>
            <div>RAS</div>
        </div>
        <div style="background: #f7f7f7; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 32px; font-weight: bold; color: #6B7280;">{len(gris)}</div>
            <div>MANQUANTS</div>
        </div>
    </div>
    {"".join([f'<div style="background: #fff5f5; border: 1px solid #fc8181; border-radius: 8px; padding: 15px; margin: 10px 0;"><h3>{r["match"]}</h3><p>{r["motif"]}</p><p>Action : {r["action"]}</p></div>' for r in rouge])}
    {"".join([f'<div style="background: #fffff0; border: 1px solid #f6e05e; border-radius: 8px; padding: 15px; margin: 10px 0;"><h3>{r["match"]}</h3><p>{r["motif"]}</p><p>Action : {r["action"]}</p></div>' for r in jaune])}
    {"".join([f'<div style="background: #f7f7f7; border: 1px solid #D1D5DB; border-radius: 8px; padding: 15px; margin: 10px 0;"><h3>{r["match"]}</h3><p>{r["motif"]}</p><p>Action : {r["action"]}</p></div>' for r in gris])}
    {"".join([f'<div style="background: #f0fff4; border: 1px solid #68d391; border-radius: 8px; padding: 15px; margin: 10px 0;"><h3>{r["match"]}</h3><p>{r["motif"]}</p></div>' for r in vert])}
    </body></html>
    """

    payload = {
        "from": "Ster-AI <onboarding@resend.dev>",
        "to": [destinataire],
        "subject": f"Ster-AI - {contexte} - {len(rouge)} prioritaire(s), {len(jaune)} a verifier, {len(gris)} manquant(s), {len(vert)} RAS",
        "html": html,
        "attachments": [
            {
                "filename": f"Recap_SterAI_{date.today().strftime('%d%m%Y')}.docx",
                "content": docx_b64
            }
        ]
    }

    response = requests.post(
        "https://api.resend.com/emails",
        headers={
            "Authorization": f"Bearer {RESEND_API_KEY}",
            "Content-Type": "application/json"
        },
        json=payload
    )

    if response.status_code != 200:
        raise Exception(f"Resend erreur {response.status_code} : {response.text}")

class RapportRequest(BaseModel):
    match: str
    rapport_arbitre: str
    rapport_delegue: str

class RecapRequest(BaseModel):
    destinataire: str
    resultats: List[dict]
    competition: Optional[str] = None
    journee: Optional[str] = None

@app.get("/")
def read_root():
    return {"message": "Ster-AI API en ligne", "status": "ok"}

@app.get("/health")
def health_check():
    return {"status": "healthy"}

@app.get("/drive/structure")
def get_drive_structure():
    try:
        service = get_drive_service()
        structure = get_structure(service)
        return {"structure": structure}
    except Exception as e:
        return {"error": str(e)}

@app.get("/drive/analyser-tout")
def analyser_tous_rapports(competition: Optional[str] = None, journee: Optional[str] = None):
    try:
        service = get_drive_service()

        if competition and journee:
            structure = get_structure(service)
            if competition not in structure:
                return {"error": f"Competition {competition} introuvable"}
            if journee not in structure[competition]['journees']:
                return {"error": f"Journee {journee} introuvable dans {competition}"}
            folder_id = structure[competition]['journees'][journee]
        else:
            folder_id = FOLDER_ID

        results = service.files().list(
            q=f"'{folder_id}' in parents and mimeType='application/pdf' and trashed=false",
            fields="files(id, name, createdTime)",
            pageSize=100
        ).execute()
        fichiers = results.get('files', [])

        arbitres = {}
        delegues = {}
        for f in fichiers:
            nom = f['name']
            if 'RA_AC_' in nom:
                id_match = nom.split('RA_AC_')[1].replace('.pdf', '')
                arbitres[id_match] = f
            elif 'RA_DP_' in nom:
                id_match = nom.split('RA_DP_')[1].replace('.pdf', '')
                delegues[id_match] = f

        tous_ids = set(arbitres.keys()) | set(delegues.keys())
        resultats = []

        for id_match in list(tous_ids)[:100]:
            rapport_arbitre = ""
            rapport_delegue = ""
            fichier_nom = f"Match {id_match}"

            if id_match in arbitres:
                rapport_arbitre = lire_pdf(service, arbitres[id_match]['id'])
                fichier_nom = arbitres[id_match]['name']
            if id_match in delegues:
                rapport_delegue = lire_pdf(service, delegues[id_match]['id'])

            if not rapport_arbitre:
                rapport_arbitre = "Rapport arbitre non disponible"
            if not rapport_delegue:
                rapport_delegue = "Rapport delegue non disponible"

            analyse = analyser_paire(id_match, rapport_arbitre, rapport_delegue)
            analyse['fichier'] = fichier_nom
            analyse['id'] = id_match
            resultats.append(analyse)

        return {"resultats": resultats, "total": len(resultats), "competition": competition, "journee": journee}
    except Exception as e:
        return {"error": str(e)}

@app.post("/envoyer-recap")
def envoyer_recap(req: RecapRequest):
    try:
        envoyer_email(req.destinataire, req.resultats, req.competition, req.journee)
        return {"message": f"Email envoye a {req.destinataire}", "total": len(req.resultats)}
    except Exception as e:
        return {"error": str(e)}

@app.post("/analyser")
def analyser_rapport(req: RapportRequest):
    prompt = f"""Tu es un expert disciplinaire de la Federation Francaise de Football.

Match : {req.match}
Rapport arbitre : {req.rapport_arbitre}
Rapport delegue : {req.rapport_delegue}

Reponds UNIQUEMENT en JSON :
{{"priorite": "rouge/jaune/vert", "motif": "...", "action": "..."}}"""

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    texte = message.content[0].text
    json_match = re.search(r'\{.*\}', texte, re.DOTALL)
    return json.loads(json_match.group())